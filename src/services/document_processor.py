"""
Document processing service for parsing and chunking various document types.
"""

import os
import re
import uuid
import asyncio
from typing import List, Optional, Dict, Any, Tuple
from datetime import datetime
import logging

# Document processing libraries
import PyPDF2
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import tiktoken

# Local imports
from ..models.document import Document, DocumentMetadata, DocumentChunk, DocumentStatus
from ..core.config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing various document types and creating text chunks."""
    
    def __init__(self):
        self.encoding = tiktoken.get_encoding("cl100k_base")
        self.supported_extensions = settings.supported_extensions
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
        self.max_file_size = settings.max_file_size_mb * 1024 * 1024  # Convert to bytes
    
    async def process_document(
        self, 
        file_path: str, 
        filename: str, 
        tags: Optional[List[str]] = None,
        custom_metadata: Optional[Dict[str, Any]] = None
    ) -> Document:
        """
        Process a document file and return a Document object with chunks.
        
        Args:
            file_path: Path to the document file
            filename: Original filename
            tags: Optional tags for the document
            custom_metadata: Optional custom metadata
            
        Returns:
            Document object with metadata and chunks
        """
        start_time = datetime.utcnow()
        
        try:
            # Validate file
            await self._validate_file(file_path, filename)
            
            # Create document metadata
            file_stats = os.stat(file_path)
            file_type = os.path.splitext(filename)[1].lower()
            
            metadata = DocumentMetadata(
                id=str(uuid.uuid4()),
                filename=filename,
                file_type=file_type,
                file_size=file_stats.st_size,
                status=DocumentStatus.PROCESSING,
                tags=tags or [],
                custom_metadata=custom_metadata
            )
            
            # Extract text content
            raw_content = await self._extract_text(file_path, file_type)
            
            # Extract document metadata from content
            extracted_metadata = await self._extract_document_metadata(raw_content, file_type)
            for key, value in extracted_metadata.items():
                if hasattr(metadata, key) and value:
                    setattr(metadata, key, value)
            
            # Create document object
            document = Document(metadata=metadata, raw_content=raw_content)
            
            # Create text chunks
            chunks = await self._create_chunks(raw_content, document.metadata.id)
            document.chunks = chunks
            document.metadata.chunk_count = len(chunks)
            
            # Update processing status
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            document.metadata.processing_time = processing_time
            document.metadata.status = DocumentStatus.COMPLETED
            
            logger.info(f"Successfully processed document {filename} with {len(chunks)} chunks")
            return document
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}")
            metadata.status = DocumentStatus.FAILED
            metadata.error_message = str(e)
            raise
    
    async def _validate_file(self, file_path: str, filename: str):
        """Validate file exists, size, and type."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_stats = os.stat(file_path)
        if file_stats.st_size > self.max_file_size:
            raise ValueError(f"File size {file_stats.st_size} exceeds maximum {self.max_file_size}")
        
        file_ext = os.path.splitext(filename)[1].lower()
        if file_ext not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    async def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text content from different file types."""
        if file_type == '.pdf':
            return await self._extract_pdf_text(file_path)
        elif file_type == '.docx':
            return await self._extract_docx_text(file_path)
        elif file_type in ['.html', '.htm']:
            return await self._extract_html_text(file_path)
        elif file_type in ['.txt', '.md']:
            return await self._extract_text_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF files."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            raise
        return text.strip()
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from Word documents."""
        try:
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise
    
    async def _extract_html_text(self, file_path: str) -> str:
        """Extract text from HTML files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
        except Exception as e:
            logger.error(f"Error extracting HTML text: {str(e)}")
            raise
    
    async def _extract_text_file(self, file_path: str) -> str:
        """Extract text from plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            logger.error(f"Error extracting text file: {str(e)}")
            raise
    
    async def _extract_document_metadata(self, content: str, file_type: str) -> Dict[str, Any]:
        """Extract metadata from document content."""
        metadata = {}
        
        # Extract title (first non-empty line or heading)
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if line and len(line) < 200:  # Reasonable title length
                metadata['title'] = line
                break
        
        # For HTML, try to extract title tag
        if file_type in ['.html', '.htm']:
            title_match = re.search(r'<title>(.*?)</title>', content, re.IGNORECASE)
            if title_match:
                metadata['title'] = title_match.group(1).strip()
        
        # Detect language (simple heuristic)
        metadata['language'] = self._detect_language(content)
        
        return metadata
    
    def _detect_language(self, text: str) -> str:
        """Simple language detection (defaults to English)."""
        # This is a placeholder - you could integrate with langdetect library
        # for more sophisticated language detection
        return "en"
    
    async def _create_chunks(self, text: str, document_id: str) -> List[DocumentChunk]:
        """Create overlapping text chunks from document content."""
        chunks = []
        
        # Clean text
        text = self._clean_text(text)
        
        # Split into sentences for better chunk boundaries
        sentences = self._split_sentences(text)
        
        current_chunk = ""
        current_tokens = 0
        chunk_index = 0
        start_char = 0
        
        for sentence in sentences:
            sentence_tokens = len(self.encoding.encode(sentence))
            
            # If adding this sentence would exceed chunk size, finalize current chunk
            if current_tokens + sentence_tokens > self.chunk_size and current_chunk:
                chunk_end_char = start_char + len(current_chunk)
                
                chunk = DocumentChunk(
                    id=f"{document_id}_chunk_{chunk_index}",
                    document_id=document_id,
                    content=current_chunk.strip(),
                    chunk_index=chunk_index,
                    start_char=start_char,
                    end_char=chunk_end_char,
                    word_count=len(current_chunk.split())
                )
                chunks.append(chunk)
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk, self.chunk_overlap)
                start_char = chunk_end_char - len(overlap_text)
                current_chunk = overlap_text + " " + sentence
                current_tokens = len(self.encoding.encode(current_chunk))
                chunk_index += 1
            else:
                # Add sentence to current chunk
                if current_chunk:
                    current_chunk += " " + sentence
                else:
                    current_chunk = sentence
                    start_char = text.find(sentence)
                
                current_tokens += sentence_tokens
        
        # Add final chunk if there's remaining content
        if current_chunk.strip():
            chunk_end_char = start_char + len(current_chunk)
            chunk = DocumentChunk(
                id=f"{document_id}_chunk_{chunk_index}",
                document_id=document_id,
                content=current_chunk.strip(),
                chunk_index=chunk_index,
                start_char=start_char,
                end_char=chunk_end_char,
                word_count=len(current_chunk.split())
            )
            chunks.append(chunk)
        
        return chunks
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\.\!\?\,\;\:\-\(\)]', ' ', text)
        
        # Remove excessive line breaks
        text = re.sub(r'\n+', '\n', text)
        
        return text.strip()
    
    def _split_sentences(self, text: str) -> List[str]:
        """Split text into sentences."""
        # Simple sentence splitting (could be improved with NLTK or spaCy)
        sentence_pattern = r'(?<=[.!?])\s+'
        sentences = re.split(sentence_pattern, text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _get_overlap_text(self, text: str, overlap_tokens: int) -> str:
        """Get the last `overlap_tokens` worth of text for chunk overlap."""
        tokens = self.encoding.encode(text)
        if len(tokens) <= overlap_tokens:
            return text
        
        overlap_token_ids = tokens[-overlap_tokens:]
        overlap_text = self.encoding.decode(overlap_token_ids)
        return overlap_text
    
    async def process_batch(
        self, 
        file_paths: List[Tuple[str, str]], 
        batch_size: Optional[int] = None
    ) -> List[Document]:
        """
        Process multiple documents in batches.
        
        Args:
            file_paths: List of (file_path, filename) tuples
            batch_size: Number of documents to process concurrently
            
        Returns:
            List of processed Document objects
        """
        if batch_size is None:
            batch_size = settings.document_processing_batch_size
        
        documents = []
        
        for i in range(0, len(file_paths), batch_size):
            batch = file_paths[i:i + batch_size]
            
            # Process batch concurrently
            tasks = [
                self.process_document(file_path, filename) 
                for file_path, filename in batch
            ]
            
            batch_results = await asyncio.gather(*tasks, return_exceptions=True)
            
            for result in batch_results:
                if isinstance(result, Exception):
                    logger.error(f"Batch processing error: {str(result)}")
                else:
                    documents.append(result)
        
        return documents 